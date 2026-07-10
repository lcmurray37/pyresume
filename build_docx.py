import json
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

def build_docx(json_path: str, output_name: str = "resume_customized_final"):
    # 1. Load Data Dynamically from JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        resume = json.load(f)

    doc = Document()

    # 2. Narrow Margins Setup (0.5 inch all around to reduce white space)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Global Style Adjustments (Arial font family)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)  # Slightly crisper size to fit text cleanly

    # Fix default spacing issues in python-docx list elements
    try:
        bullet_style = doc.styles['List Bullet']
        bullet_font = bullet_style.font
        bullet_font.name = 'Arial'
        bullet_font.size = Pt(10)
    except KeyError:
        pass

    # 3. Header Block (Name & Contacts)
    title = doc.add_paragraph()
    title.alignment = 1  # Centered
    title_run = title.add_run(resume.get("name", "Name Missing"))
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)

    c = resume.get("contact", {})
    contact_p = doc.add_paragraph()
    contact_p.alignment = 1
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(12) # Tightened down header buffer
    
    c_run = contact_p.add_run(f"{c.get('email', '')}  |  {c.get('phone', '')}  |  {c.get('github', '')}")
    c_run.font.name = 'Arial'
    c_run.font.size = Pt(10)

    # --- Helper function for clean, tight Section Titles ---
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10) # Reduced from 12
        p.paragraph_format.space_after = Pt(2)  # Reduced from 4
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        p.paragraph_format.keep_with_next = True

    # Helper function for tight bullet formatting
    def add_bullet_point(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2) # Keeps list items physically tight
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    # 4. Professional Summary Section
    summary = resume.get("summary", [])
    if summary:
        add_section_heading("Professional Summary")
        for line in summary:
            add_bullet_point(line)

    # 5. Experience Section
    experience = resume.get("experience", [])
    if experience:
        add_section_heading("Experience")
        for job in experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            
            comp = job.get("company", "")
            title_str = job.get("title", "")
            dates = job.get("dates", "")
            
            run_title = p.add_run(f"{comp} | {title_str} ")
            run_title.bold = True
            run_title.font.name = 'Arial'
            
            run_dates = p.add_run(f"({dates})")
            run_dates.italic = True
            run_dates.font.name = 'Arial'
            
            # Subsections (e.g., Project Management, Advanced Analytics)
            for section_title, bullets in job.get("sections", {}).items():
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(2)
                sp.paragraph_format.space_after = Pt(1)
                
                s_run = sp.add_run(section_title)
                s_run.bold = True
                s_run.font.name = 'Arial'
                s_run.font.size = Pt(10)
                
                # Achievement Bullets
                for bullet in bullets:
                    add_bullet_point(bullet)

    # 6. Education Section
    education = resume.get("education", [])
    if education:
        add_section_heading("Education")
        for edu in education:
            add_bullet_point(edu)

    # 7. Core Competencies Section
    skills = resume.get("skills", {})
    if skills:
        add_section_heading("Core Competencies")
        for category, skill_list in skills.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            
            cat_run = p.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.font.name = 'Arial'
            cat_run.font.size = Pt(10)
            
            sk_run = p.add_run(skill_list)
            sk_run.font.name = 'Arial'
            sk_run.font.size = Pt(10)

    # 8. Save Document
    output_dir = Path("build")
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / f"{output_name}.docx"
    doc.save(output_path)
    print(f"✅ Successfully created narrow-margin Arial Word Document: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert tailored JSON resume metadata to a tight Arial DOCX file.")
    parser.add_argument("-i", "--input", required=True, help="Path to the tailored JSON resume file.")
    parser.add_argument("-o", "--output", default="resume_customized_final", help="Output Word document name.")
    
    args = parser.parse_args()
    build_docx(args.input, args.output)